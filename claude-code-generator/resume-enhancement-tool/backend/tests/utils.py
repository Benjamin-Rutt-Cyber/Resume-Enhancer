"""
Test utilities for creating test data and helper functions.

This module provides utility functions for:
- Creating test PDF and DOCX files
- Creating test database records
- Common test data generation
"""

import io
from pathlib import Path
from typing import Dict, Any, Optional
from uuid import UUID, uuid4

from sqlalchemy.orm import Session

from app.models import Resume, Job, Enhancement

# Optional dependencies for PDF/DOCX creation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


# ============================================================================
# PDF Generation Utilities
# ============================================================================

def create_test_pdf(content: str, file_path: Optional[Path] = None) -> Path:
    """
    Create a test PDF file with the given content.

    Args:
        content: Text content to include in the PDF
        file_path: Optional path where to save the PDF.
                  If not provided, creates in a temp directory.

    Returns:
        Path: Path to the created PDF file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not REPORTLAB_AVAILABLE:
        # Fallback: create a simple text file with .pdf extension for testing
        # Real tests should have reportlab installed
        file_path.write_text(content, encoding='utf-8')
        return file_path

    # Create PDF using reportlab
    c = canvas.Canvas(str(file_path), pagesize=letter)
    width, height = letter

    # Add text to PDF
    text_object = c.beginText(50, height - 50)
    text_object.setFont("Helvetica", 12)

    # Split content into lines and add to PDF
    lines = content.split('\n')
    for line in lines:
        text_object.textLine(line.strip())

    c.drawText(text_object)
    c.showPage()
    c.save()

    return file_path


def create_test_pdf_multipage(content: str, pages: int = 3, file_path: Optional[Path] = None) -> Path:
    """
    Create a multi-page test PDF file.

    Args:
        content: Text content to include on each page
        pages: Number of pages to create
        file_path: Optional path where to save the PDF

    Returns:
        Path: Path to the created PDF file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not REPORTLAB_AVAILABLE:
        # Fallback: create extended content
        multipage_content = "\n\n".join([f"Page {i+1}\n{content}" for i in range(pages)])
        file_path.write_text(multipage_content, encoding='utf-8')
        return file_path

    c = canvas.Canvas(str(file_path), pagesize=letter)
    width, height = letter

    for page_num in range(1, pages + 1):
        text_object = c.beginText(50, height - 50)
        text_object.setFont("Helvetica", 12)
        text_object.textLine(f"Page {page_num}")
        text_object.textLine("")

        lines = content.split('\n')
        for line in lines:
            text_object.textLine(line.strip())

        c.drawText(text_object)
        c.showPage()

    c.save()

    return file_path


def create_empty_pdf(file_path: Optional[Path] = None) -> Path:
    """
    Create an empty PDF file (for testing error handling).

    Args:
        file_path: Optional path where to save the PDF

    Returns:
        Path: Path to the created PDF file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not REPORTLAB_AVAILABLE:
        # Fallback: create empty text file
        file_path.write_text("", encoding='utf-8')
        return file_path

    c = canvas.Canvas(str(file_path), pagesize=letter)
    c.showPage()
    c.save()

    return file_path


# ============================================================================
# DOCX Generation Utilities
# ============================================================================

def create_test_docx(content: str, file_path: Optional[Path] = None) -> Path:
    """
    Create a test DOCX file with the given content.

    Args:
        content: Text content to include in the DOCX
        file_path: Optional path where to save the DOCX

    Returns:
        Path: Path to the created DOCX file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not PYTHON_DOCX_AVAILABLE:
        # Fallback: create text file for testing
        file_path.write_text(content, encoding='utf-8')
        return file_path

    doc = Document()

    # Add content as paragraphs
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.save(str(file_path))

    return file_path


def create_test_docx_with_table(content: str, file_path: Optional[Path] = None) -> Path:
    """
    Create a test DOCX file with a table (common in resumes).

    Args:
        content: Text content to include in the DOCX
        file_path: Optional path where to save the DOCX

    Returns:
        Path: Path to the created DOCX file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not PYTHON_DOCX_AVAILABLE:
        # Fallback: create text with table-like structure
        table_content = """Resume
Name: John Doe
Email: john@example.com
Phone: 555-1234

""" + content
        file_path.write_text(table_content, encoding='utf-8')
        return file_path

    doc = Document()

    # Add some text
    doc.add_heading('Resume', 0)

    # Add a table (2 columns x 3 rows)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'

    # Fill table with sample data
    cells = [
        ['Name:', 'John Doe'],
        ['Email:', 'john@example.com'],
        ['Phone:', '555-1234'],
    ]

    for i, row_data in enumerate(cells):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]

    # Add content
    doc.add_paragraph(content)

    doc.save(str(file_path))

    return file_path


def create_empty_docx(file_path: Optional[Path] = None) -> Path:
    """
    Create an empty DOCX file (for testing error handling).

    Args:
        file_path: Optional path where to save the DOCX

    Returns:
        Path: Path to the created DOCX file
    """
    if file_path is None:
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        file_path = Path(temp_file.name)
        temp_file.close()

    if not PYTHON_DOCX_AVAILABLE:
        # Fallback: create empty text file
        file_path.write_text("", encoding='utf-8')
        return file_path

    doc = Document()
    doc.save(str(file_path))

    return file_path


# ============================================================================
# Database Model Utilities
# ============================================================================

def create_test_resume_in_db(
    db: Session,
    filename: str = "test_resume.pdf",
    word_count: int = 150,
    **kwargs
) -> Resume:
    """
    Create a test resume record in the database.

    Args:
        db: Database session
        filename: Resume filename
        word_count: Number of words in resume
        **kwargs: Additional fields to override

    Returns:
        Resume: Created resume model instance
    """
    resume_id = kwargs.pop('id', uuid4())

    resume = Resume(
        id=resume_id,
        filename=filename,
        original_format=kwargs.pop('original_format', 'pdf'),
        file_path=kwargs.pop('file_path', f"/workspace/resumes/original/{resume_id}/source.pdf"),
        extracted_text_path=kwargs.pop('extracted_text_path', f"/workspace/resumes/original/{resume_id}/extracted.txt"),
        file_size_bytes=kwargs.pop('file_size_bytes', 50000),
        word_count=word_count,
        **kwargs
    )

    db.add(resume)
    db.commit()
    db.refresh(resume)

    return resume


def create_test_job_in_db(
    db: Session,
    title: str = "Software Engineer",
    company: str = "Tech Company",
    **kwargs
) -> Job:
    """
    Create a test job record in the database.

    Args:
        db: Database session
        title: Job title
        company: Company name
        **kwargs: Additional fields to override

    Returns:
        Job: Created job model instance
    """
    job = Job(
        id=kwargs.pop('id', uuid4()),
        title=title,
        company=company,
        description=kwargs.pop('description', "We are looking for a talented software engineer..."),
        source=kwargs.pop('source', None),
        **kwargs
    )

    db.add(job)
    db.commit()
    db.refresh(job)

    return job


def create_test_enhancement_in_db(
    db: Session,
    resume_id: UUID,
    job_id: Optional[UUID] = None,
    enhancement_type: str = "job_tailoring",
    **kwargs
) -> Enhancement:
    """
    Create a test enhancement record in the database.

    Args:
        db: Database session
        resume_id: ID of the resume to enhance
        job_id: Optional ID of the job (required for job_tailoring)
        enhancement_type: Type of enhancement (job_tailoring or industry_revamp)
        **kwargs: Additional fields to override

    Returns:
        Enhancement: Created enhancement model instance
    """
    enhancement = Enhancement(
        id=kwargs.pop('id', uuid4()),
        resume_id=resume_id,
        job_id=job_id,
        enhancement_type=enhancement_type,
        status=kwargs.pop('status', 'pending'),
        industry=kwargs.pop('industry', None),
        **kwargs
    )

    db.add(enhancement)
    db.commit()
    db.refresh(enhancement)

    return enhancement


# ============================================================================
# Sample Text Data
# ============================================================================

SAMPLE_RESUME_SHORT = """
John Doe
Software Engineer
john@example.com
"""

SAMPLE_RESUME_VALID = """
John Doe
Software Engineer
john.doe@email.com | (555) 123-4567

SUMMARY
Experienced software engineer with 5+ years of expertise in full-stack development.
Proven track record of delivering scalable applications.

EXPERIENCE
Senior Software Engineer | Tech Company | 2020 - Present
- Led development of microservices architecture
- Reduced API response time by 40%

EDUCATION
B.S. Computer Science | University | 2018

SKILLS
Python, JavaScript, TypeScript, Docker, AWS
"""

SAMPLE_RESUME_LONG = """
John Doe
Senior Software Engineer
john.doe@email.com | (555) 123-4567 | LinkedIn: johndoe | GitHub: johndoe

PROFESSIONAL SUMMARY
Highly accomplished software engineer with over 8 years of experience in full-stack web development,
cloud architecture, and team leadership. Expert in Python, JavaScript, and modern DevOps practices.
Proven track record of designing and implementing scalable systems serving millions of users.

TECHNICAL SKILLS
Languages: Python, JavaScript, TypeScript, Java, Go, SQL
Frameworks: FastAPI, Django, React, Node.js, Express, Spring Boot
Cloud: AWS (EC2, S3, Lambda, RDS, CloudFormation), Azure, Google Cloud
DevOps: Docker, Kubernetes, Jenkins, GitHub Actions, Terraform
Databases: PostgreSQL, MySQL, MongoDB, Redis, Elasticsearch
Tools: Git, Linux, Nginx, Apache, Kafka

PROFESSIONAL EXPERIENCE

Senior Software Engineer | Tech Corp Inc. | San Francisco, CA | 2020 - Present
- Architected and implemented microservices-based platform serving 2M+ active users
- Led team of 7 engineers in development of real-time data processing pipeline
- Reduced infrastructure costs by 35% through optimization of AWS resources
- Implemented CI/CD pipeline reducing deployment time from 2 hours to 15 minutes
- Mentored 10+ junior developers on best practices and code review

Software Engineer | StartUp Co. | New York, NY | 2018 - 2020
- Built RESTful APIs using Python FastAPI handling 10K requests/second
- Developed React-based admin dashboard reducing customer support time by 40%
- Implemented OAuth2 authentication and role-based access control
- Collaborated with product team on feature prioritization and sprint planning
- Achieved 95% test coverage across all microservices

Junior Developer | Software Solutions Ltd. | Boston, MA | 2016 - 2018
- Developed full-stack web applications using Django and React
- Participated in agile development process with 2-week sprints
- Fixed over 200 bugs and implemented 50+ new features
- Improved application performance by optimizing database queries

EDUCATION
Master of Science in Computer Science | Stanford University | 2016
Bachelor of Science in Computer Science | MIT | 2014

CERTIFICATIONS
- AWS Certified Solutions Architect - Professional
- Certified Kubernetes Administrator (CKA)
- Google Cloud Professional Cloud Architect

PROJECTS
Open Source Contributions:
- Contributor to FastAPI framework (500+ GitHub stars)
- Maintainer of popular Python CLI library (1K+ downloads/month)

Personal Projects:
- Built SaaS platform for team collaboration (500+ active users)
- Created developer tool that reduces deployment time by 60%
"""

SAMPLE_JOB_DESCRIPTION = """
Desktop Support Engineer

Total IT Global is seeking an experienced Desktop Support Engineer.

Responsibilities:
- Provide technical support to end users
- Troubleshoot hardware and software issues
- Install and configure systems
- Maintain IT documentation
- Support remote users via VPN

Requirements:
- 2+ years IT support experience
- Windows 10/11 expertise
- Active Directory knowledge
- Excellent communication skills
- CompTIA A+ preferred

Skills:
Windows, MacOS, Office 365, Remote Desktop, VPN, Networking basics
"""


# ============================================================================
# Cleanup Utilities
# ============================================================================

def cleanup_workspace(workspace_path: Path):
    """
    Clean up test workspace directory.

    Args:
        workspace_path: Path to workspace directory to clean up
    """
    import shutil

    if workspace_path.exists():
        shutil.rmtree(workspace_path)
