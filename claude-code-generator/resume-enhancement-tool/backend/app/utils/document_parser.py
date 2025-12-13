"""Document parser for PDF and DOCX resumes."""

from pathlib import Path
from typing import Dict, Optional
import pdfplumber
from pypdf import PdfReader
from docx import Document


class DocumentParser:
    """Parse PDF and DOCX documents and extract text content."""

    def parse_file(self, file_path: Path) -> Dict[str, any]:
        """
        Parse a document file and extract its text content.

        Args:
            file_path: Path to the PDF or DOCX file

        Returns:
            Dictionary with extracted text, format, and metadata

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _parse_pdf(self, file_path: Path) -> Dict[str, any]:
        """
        Parse a PDF file using pdfplumber (with pypdf fallback).

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Try pdfplumber first (better formatting and layout awareness)
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)

                full_text = "\n\n".join(pages_text)

                return {
                    "text": full_text,
                    "format": "pdf",
                    "pages": len(pdf.pages),
                    "parser": "pdfplumber",
                    "success": True,
                }
        except Exception as e:
            # Fallback to pypdf if pdfplumber fails
            try:
                reader = PdfReader(file_path)
                pages_text = []

                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)

                full_text = "\n\n".join(pages_text)

                return {
                    "text": full_text,
                    "format": "pdf",
                    "pages": len(reader.pages),
                    "parser": "pypdf",
                    "success": True,
                }
            except Exception as fallback_error:
                return {
                    "text": "",
                    "format": "pdf",
                    "pages": 0,
                    "parser": "failed",
                    "success": False,
                    "error": str(fallback_error),
                }

    def _parse_docx(self, file_path: Path) -> Dict[str, any]:
        """
        Parse a DOCX file and extract text from paragraphs and tables.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract text from tables (resumes often use tables for formatting)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

            # Combine paragraphs and table text
            all_text = paragraphs + table_text
            full_text = "\n".join(all_text)

            return {
                "text": full_text,
                "format": "docx",
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "parser": "python-docx",
                "success": True,
            }
        except Exception as e:
            return {
                "text": "",
                "format": "docx",
                "paragraphs": 0,
                "tables": 0,
                "parser": "python-docx",
                "success": False,
                "error": str(e),
            }

    def validate_content(self, text: str, min_words: int = 50) -> Dict[str, any]:
        """
        Validate that extracted text contains meaningful content.

        Args:
            text: Extracted text
            min_words: Minimum number of words expected

        Returns:
            Dictionary with validation results
        """
        if not text or not text.strip():
            return {
                "valid": False,
                "reason": "No text extracted",
                "word_count": 0,
            }

        words = text.split()
        word_count = len(words)

        if word_count < min_words:
            return {
                "valid": False,
                "reason": f"Too few words extracted (got {word_count}, expected at least {min_words})",
                "word_count": word_count,
            }

        # Check for common resume indicators
        resume_keywords = [
            "experience",
            "education",
            "skills",
            "work",
            "employment",
            "university",
            "degree",
            "certification",
        ]

        text_lower = text.lower()
        found_keywords = [kw for kw in resume_keywords if kw in text_lower]

        if not found_keywords:
            return {
                "valid": True,  # Still valid, just a warning
                "reason": "No typical resume keywords found - may not be a resume",
                "word_count": word_count,
                "warning": True,
            }

        return {
            "valid": True,
            "reason": "Content looks valid",
            "word_count": word_count,
            "found_keywords": found_keywords,
        }
