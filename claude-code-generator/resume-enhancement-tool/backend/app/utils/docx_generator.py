"""Convert markdown resume to styled DOCX."""

from pathlib import Path
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


class DOCXGenerator:
    """Generate styled DOCX from markdown resume."""

    def __init__(self):
        """Initialize DOCX generator with default styling."""
        self.heading_sizes = {
            1: 18,  # Main name/title
            2: 14,  # Section headers
            3: 12   # Subsection headers
        }
        self.heading_colors = {
            1: RGBColor(0, 51, 102),    # Dark blue
            2: RGBColor(0, 102, 204),   # Medium blue
            3: RGBColor(51, 51, 51),    # Dark gray
        }

    def markdown_to_docx(self, md_path: Path, docx_path: Path) -> None:
        """Convert markdown resume to styled DOCX.

        Args:
            md_path: Path to input markdown file
            docx_path: Path to output DOCX file

        Raises:
            FileNotFoundError: If markdown file doesn't exist
            Exception: If conversion fails
        """

        if not md_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_path}")

        # Read markdown content
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Create new document
        doc = Document()

        # Set page margins
        self._set_margins(doc)

        # Parse markdown line by line
        self._parse_and_add_content(doc, md_content)

        # Save document
        doc.save(str(docx_path))

    def _set_margins(self, doc: Document) -> None:
        """Set document margins to 0.75 inches.

        Args:
            doc: Document object to modify
        """
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _parse_and_add_content(self, doc: Document, md_content: str) -> None:
        """Parse markdown and add formatted content to document.

        Args:
            doc: Document object to add content to
            md_content: Markdown text to parse
        """

        lines = md_content.split('\n')
        in_list = False
        skip_next = False

        for i, line in enumerate(lines):
            line = line.rstrip()

            # Skip this line if flagged
            if skip_next:
                skip_next = False
                continue

            # Empty line - add spacing
            if not line:
                # Close list if we were in one
                if in_list:
                    in_list = False
                doc.add_paragraph()
                continue

            # Horizontal rule (---)
            if line.strip() == '---':
                # Add a subtle separator paragraph
                p = doc.add_paragraph()
                p.add_run('_' * 80)
                run = p.runs[0]
                run.font.color.rgb = RGBColor(200, 200, 200)
                run.font.size = Pt(6)
                continue

            # Heading 1 (# Name)
            if line.startswith('# '):
                text = line[2:].strip()
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(self.heading_sizes[1])
                    run.font.color.rgb = self.heading_colors[1]
                    run.font.bold = True
                continue

            # Heading 2 (## Section)
            if line.startswith('## '):
                text = line[3:].strip()
                p = doc.add_heading(text, level=2)
                for run in p.runs:
                    run.font.size = Pt(self.heading_sizes[2])
                    run.font.color.rgb = self.heading_colors[2]
                    run.font.bold = True
                continue

            # Heading 3 (### Subsection)
            if line.startswith('### '):
                text = line[4:].strip()
                p = doc.add_heading(text, level=3)
                for run in p.runs:
                    run.font.size = Pt(self.heading_sizes[3])
                    run.font.bold = True
                continue

            # Bullet list (- item or * item)
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
                in_list = True
                continue

            # Numbered list (1. item, 2. item, etc.)
            if re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, text)
                in_list = True
                continue

            # Regular paragraph
            if in_list:
                in_list = False

            p = doc.add_paragraph()
            self._add_formatted_text(p, line)

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text with inline formatting (bold, italic, links) to paragraph.

        Args:
            paragraph: Paragraph object to add text to
            text: Text with markdown inline formatting

        Supports:
            - **bold text**
            - *italic text*
            - [link text](url) - renders as blue underlined text
        """

        # Process bold, italic, and links
        # Pattern priority: bold first (two asterisks), then links, then italic

        remaining = text
        while remaining:
            # Check for bold (**text**)
            bold_match = re.search(r'\*\*(.+?)\*\*', remaining)
            # Check for italic (*text* but not **)
            italic_match = re.search(r'(?<!\*)\*([^*]+?)\*(?!\*)', remaining)
            # Check for links ([text](url))
            link_match = re.search(r'\[([^\]]+)\]\(([^\)]+)\)', remaining)

            # Find which match comes first
            matches = []
            if bold_match:
                matches.append(('bold', bold_match.start(), bold_match))
            if italic_match:
                matches.append(('italic', italic_match.start(), italic_match))
            if link_match:
                matches.append(('link', link_match.start(), link_match))

            if not matches:
                # No more formatting, add remaining text
                if remaining:
                    paragraph.add_run(remaining)
                break

            # Sort by position and take the first
            matches.sort(key=lambda x: x[1])
            format_type, pos, match = matches[0]

            # Add text before the match
            if pos > 0:
                paragraph.add_run(remaining[:pos])

            # Add formatted text
            if format_type == 'bold':
                run = paragraph.add_run(match.group(1))
                run.bold = True
                remaining = remaining[match.end():]
            elif format_type == 'italic':
                run = paragraph.add_run(match.group(1))
                run.italic = True
                remaining = remaining[match.end():]
            elif format_type == 'link':
                # Render link as blue underlined text
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.underline = True
                remaining = remaining[match.end():]

    def create_from_text(self, text: str, output_path: Path, title: Optional[str] = None) -> None:
        """Create a DOCX document directly from text (without markdown file).

        Args:
            text: Markdown-formatted text
            output_path: Path to save DOCX file
            title: Optional document title
        """

        doc = Document()
        self._set_margins(doc)

        if title:
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.font.bold = True

        self._parse_and_add_content(doc, text)
        doc.save(str(output_path))
